import uuid
from io import BytesIO
from datetime import datetime
from docx import Document
from fastapi import UploadFile, HTTPException

from app.storage.documents import documents_store


class DocumentService:
    @staticmethod
    async def extract_text_from_docx(file_content: bytes) -> str:
        try:
            doc = Document(BytesIO(file_content))
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            full_text.append(para.text)

            return "\n".join(full_text)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")

    @staticmethod
    async def upload_document(file: UploadFile) -> str:
        if not file.filename.endswith(".docx"):
            raise HTTPException(
                status_code=400, detail="Only .docx files are supported"
            )

        content = await file.read()
        try:
            text = await DocumentService.extract_text_from_docx(content)
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))

        file_id = str(uuid.uuid4())

        documents_store[file_id] = {
            "filename": file.filename,
            "text": text,
            "upload_time": datetime.now(),
        }

        return file_id
